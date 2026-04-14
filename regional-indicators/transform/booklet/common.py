"""
transform/booklet/common.py
---------------------------
Shared constants, styles, and helpers for booklet generation.
"""
from pathlib import Path
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ──────────────────────────────────────────────────────────────────────
REPO_ROOT   = Path(__file__).parent.parent.parent  # regional-indicators/
RESOURCES   = REPO_ROOT / "resources" / "images"
SHARED_IMG  = RESOURCES / "shared"
TOWNS_IMG   = RESOURCES / "towns"
BOOKLETS    = REPO_ROOT / "booklets"
OUTPUT_DIR  = REPO_ROOT / "output"

# ── Brand colours ──────────────────────────────────────────────────────────────
PURPLE      = RGBColor(0x51, 0x24, 0x7A)   # UQ purple
TEAL        = RGBColor(0x00, 0xA9, 0xCE)   # UQ teal
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREY_DARK   = RGBColor(0x44, 0x44, 0x44)
GREY_MED    = RGBColor(0xCC, 0xCC, 0xCC)
GREY_LIGHT  = RGBColor(0xF2, 0xF2, 0xF2)

# Hex strings for XML-level shading
HEX_PURPLE      = "51247A"
HEX_TEAL        = "00A9CE"
HEX_WHITE       = "FFFFFF"
HEX_GREY_LIGHT  = "F2F2F2"
HEX_GREY_MED    = "CCCCCC"

# ── Page dimensions (A4 portrait, in EMU) ──────────────────────────────────────
# 1 inch = 914400 EMU; A4 = 210mm x 297mm
PAGE_W_MM   = 210
PAGE_H_MM   = 297
MARGIN_MM   = 20
CONTENT_W_MM = PAGE_W_MM - 2 * MARGIN_MM   # 170mm

PAGE_W_EMU  = int(PAGE_W_MM  * 914400 / 25.4)
PAGE_H_EMU  = int(PAGE_H_MM  * 914400 / 25.4)
MARGIN_EMU  = int(MARGIN_MM  * 914400 / 25.4)
CONTENT_W_EMU = int(CONTENT_W_MM * 914400 / 25.4)

# ── Typography helpers ─────────────────────────────────────────────────────────

def set_run_font(run, size_pt, bold=False, italic=False, colour=None):
    run.font.name   = "Arial"
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour

def add_para(doc, text, size_pt, bold=False, italic=False, colour=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size_pt, bold=bold, italic=italic, colour=colour or GREY_DARK)
    return p

def add_spacer(doc, points=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(points)
    return p

# ── Table helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour):
    """Set cell background shading colour.
    In tcPr schema order: tcW → tcBorders → shd → ...
    We insert shd after tcBorders if present, else after tcW, else at end.
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd to avoid duplicates
    for ex in tcPr.findall(qn("w:shd")):
        tcPr.remove(ex)
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    # Insert after tcBorders if present, else after tcW, else append
    tcb = tcPr.find(qn("w:tcBorders"))
    tcw = tcPr.find(qn("w:tcW"))
    if tcb is not None:
        tcb.addnext(shd)
    elif tcw is not None:
        tcw.addnext(shd)
    else:
        tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders. Pass None for a side to remove it, or a dict with
    style/size/color keys to set it."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing tcBorders element
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    # OOXML strict uses w:start/w:end; transitional uses w:left/w:right.
    # python-docx targets transitional, so use left/right.
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if val is None:
            el.set(qn("w:val"), "none")
        else:
            el.set(qn("w:val"),   val.get("style", "single"))
            el.set(qn("w:sz"),    str(val.get("size", 4)))
            el.set(qn("w:color"), val.get("color", "000000"))
        tcBorders.append(el)
    # tcBorders must come after tcW in tcPr schema
    tcW_el = tcPr.find(qn("w:tcW"))
    if tcW_el is not None:
        tcW_el.addnext(tcBorders)
    else:
        tcPr.append(tcBorders)

def no_borders(cell):
    set_cell_borders(cell)  # all None = none

def set_table_width(table, width_emu):
    """Replace (not append) the tblW element so there is never a duplicate."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Remove all existing tblW elements
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(int(width_emu / 914400 * 1440)))
    tblW.set(qn("w:type"), "dxa")
    # tblW must follow tblStyle but precede jc/borders in schema
    tbl_style = tblPr.find(qn("w:tblStyle"))
    if tbl_style is not None:
        tbl_style.addnext(tblW)
    else:
        tblPr.insert(0, tblW)

def set_col_width(cell, width_twips):
    """Set cell width in twips (1440 twips = 1 inch). Replaces any existing tcW."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_twips)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.insert(0, tcW)  # tcW must be first child in tcPr

# ── Image helpers ──────────────────────────────────────────────────────────────

def add_image_centered(doc, path, width_mm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_mm / 25.4))
    return p

# ── Section header ─────────────────────────────────────────────────────────────

def add_section_header(doc, title, subtitle=None):
    """Purple banner heading used on each data page."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    # Purple background via shading on paragraph
    p.paragraph_format.left_indent  = Pt(10)
    p.paragraph_format.right_indent = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  HEX_PURPLE)
    # shd must precede spacing and ind — insert at position 0 (before pStyle would be
    # wrong, but we never set pStyle so pos 0 is the first child = correct)
    pPr.insert(0, shd)
    run = p.add_run(title)
    set_run_font(run, 16, bold=True, colour=WHITE)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(8)
        p2.paragraph_format.left_indent  = Pt(10)
        p2.paragraph_format.right_indent = Pt(10)
        pPr2 = p2._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"),   "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"),  HEX_PURPLE)
        pPr2.insert(0, shd2)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, 10, italic=True, colour=WHITE)

# ── Footer line ────────────────────────────────────────────────────────────────

def add_footer_line(doc, town_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "4")
    top.set(qn("w:color"), HEX_PURPLE)
    top.set(qn("w:space"), "4")
    pBdr.append(top)
    pPr.insert(0, pBdr)  # pBdr must precede spacing/jc in pPr schema
    run = p.add_run(f"This information has been compiled for use in consultation with the {town_name} community")
    set_run_font(run, 8, italic=True, colour=PURPLE)