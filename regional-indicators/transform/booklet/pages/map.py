"""
transform/booklet/pages/map.py
-------------------------------
Page 3: SA2 boundary map page.

Layout:
  - Page heading (not printed — orientation context only)
  - Map image filling content width
  - Bold caption: "{sa2_name} statistical area 2 (SA2) boundary"
  - Source URL (plain text)
  - Centred footer: "This information has been compiled for use in
    consultation with the {town} community"

Map PDF source:
  QLD towns: exact URL stored in towns.toml as sa2_map_url
  Fetched, rasterised at 150 DPI, cropped to main map area,
  cached as resources/images/towns/{slug}/sa2_map.png.
  Delete the cached PNG to force a refresh.
"""
from __future__ import annotations

import urllib.request
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
_BOOKLET_DIR = Path(__file__).resolve().parent.parent
if str(_BOOKLET_DIR) not in sys.path:
    sys.path.insert(0, str(_BOOKLET_DIR))

from common import (
    TOWNS_IMG, PURPLE, GREY_DARK, HEX_GREY_LIGHT, HEX_PURPLE,
    CONTENT_W_MM,
    set_run_font, add_footer_line,
    set_cell_bg, no_borders,
)

MAP_DPI = 150


# ── Fetch + rasterise ─────────────────────────────────────────────────────────

def _fetch_and_render(pdf_url: str, out_png: Path) -> bool:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        print("    WARNING: pypdfium2 not installed. Run: pip install pypdfium2")
        return False

    pdf_tmp = out_png.parent / "_sa2_tmp.pdf"
    out_png.parent.mkdir(parents=True, exist_ok=True)

    print(f"    Fetching: {pdf_url}")
    try:
        req = urllib.request.Request(pdf_url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            pdf_tmp.write_bytes(resp.read())
    except Exception as e:
        print(f"    WARNING: Could not fetch map PDF: {e}")
        pdf_tmp.unlink(missing_ok=True)
        return False

    try:
        pdf    = pdfium.PdfDocument(str(pdf_tmp))
        page   = pdf[0]
        bitmap = page.render(scale=MAP_DPI / 72)
        img    = bitmap.to_pil().convert("RGB")
        w, h   = img.size
        # QGSO PDF layout (A4 landscape):
        #   top ~8%  : title bar  → strip (heading is in the docx)
        #   left ~65%: main map
        #   right 35%: QLD inset + legend + credits → exclude
        title_h  = int(h * 0.08)
        map_crop = img.crop((0, title_h, int(w * 0.67), h))
        map_crop.save(str(out_png))
        pdf_tmp.unlink(missing_ok=True)
        print(f"    Cached:  {out_png.name} ({map_crop.size[0]}x{map_crop.size[1]}px)")
        return True
    except Exception as e:
        print(f"    WARNING: Could not rasterise map PDF: {e}")
        pdf_tmp.unlink(missing_ok=True)
        return False


def get_map_png(town_cfg: dict) -> Path | None:
    slug    = town_cfg["slug"]
    pdf_url = town_cfg.get("sa2_map_url", "")
    out_png = TOWNS_IMG / slug / "sa2_map.png"

    if out_png.exists():
        print(f"    Using cached map: {out_png.name}")
        return out_png

    if not pdf_url:
        print(f"    WARNING: No sa2_map_url in towns.toml for {slug}")
        return None

    return out_png if _fetch_and_render(pdf_url, out_png) else None


# ── Page builder ──────────────────────────────────────────────────────────────

def build_map_page(doc: Document, town_cfg: dict):
    town_name = town_cfg["name"]
    sa2_name  = town_cfg.get("sa2_name", town_name)
    sa2_code  = town_cfg.get("sa2_code", "")
    pdf_url   = town_cfg.get("sa2_map_url", "")

    # ── Page title (top of landscape page) ────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(8)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        title_p.add_run(
            "UQ Research Project: Cumulative social and economic impacts "
            "of CSG development in Western Downs"
        ),
        10, colour=GREY_DARK
    )

    # ── Map image or placeholder ───────────────────────────────────────────────
    map_png = get_map_png(town_cfg)

    if map_png and map_png.exists():
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        img_p.paragraph_format.space_before = Pt(0)
        img_p.paragraph_format.space_after  = Pt(6)
        img_p.add_run().add_picture(
            str(map_png), width=Inches(CONTENT_W_MM / 25.4))
    else:
        _placeholder(doc, town_name, sa2_code, pdf_url)

    # ── Caption (bold) ─────────────────────────────────────────────────────────
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(2)
    set_run_font(
        cap.add_run(f"{sa2_name} statistical area 2 (SA2) boundary"),
        9, bold=True, colour=GREY_DARK
    )

    # ── Source URL ─────────────────────────────────────────────────────────────
    src = doc.add_paragraph()
    src.paragraph_format.space_before = Pt(0)
    src.paragraph_format.space_after  = Pt(0)
    set_run_font(src.add_run("Source: "), 8, colour=GREY_DARK)
    set_run_font(src.add_run(pdf_url or "See towns.toml for sa2_map_url"), 8, colour=PURPLE)

    # ── Footer ─────────────────────────────────────────────────────────────────
    add_footer_line(doc, town_name)


# ── Placeholder ───────────────────────────────────────────────────────────────

def _placeholder(doc, town_name, sa2_code, pdf_url):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, HEX_GREY_LIGHT)

    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), HEX_PURPLE)
        tcBorders.append(el)
    tcW_el = tcPr.find(qn("w:tcW"))
    if tcW_el is not None:
        tcW_el.addnext(tcBorders)
    else:
        tcPr.append(tcBorders)

    tr   = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(110 / 25.4 * 1440)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)

    items = [
        (f"[ SA2 MAP — {town_name.upper()} ]",    True,  12, 36),
        (f"SA2 Code: {sa2_code}",                  False,  9,  6),
        ("Map could not be fetched automatically.", True,   8, 10),
        ("Download from:",                          False,  8,  4),
        (pdf_url or "Add sa2_map_url to towns.toml", False, 8, 2),
        (f"Save as: resources/images/towns/{town_name.lower()}/sa2_map.png", False, 8, 2),
        (f"Then re-run: python transform/booklet/make_booklet.py --town {town_name}", False, 8, 2),
    ]
    first = True
    for text, bold, size, top_space in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(top_space)
        p.paragraph_format.space_after  = Pt(0)
        set_run_font(p.add_run(text), size, bold=bold,
                     colour=PURPLE if bold else GREY_DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)