"""
transform/booklet/pages/cover.py
---------------------------------
Page 1: Cover page (Portrait A4).

Standalone test:
    python transform/booklet/pages/cover.py --town Chinchilla
"""
from __future__ import annotations

import argparse
import sys
import tomllib
import zipfile
import shutil
from datetime import datetime
from pathlib import Path

# ── Path setup: must happen before any local imports ──────────────────────────
# cover.py lives at: booklet/pages/cover.py
# common.py lives at: booklet/common.py
# So we add booklet/ (parent of pages/) to sys.path
_PAGES_DIR   = Path(__file__).resolve().parent
_BOOKLET_DIR = _PAGES_DIR.parent
_REPO_ROOT   = _BOOKLET_DIR.parent.parent
if str(_BOOKLET_DIR) not in sys.path:
    sys.path.insert(0, str(_BOOKLET_DIR))
# ─────────────────────────────────────────────────────────────────────────────

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from common import (
    SHARED_IMG, TOWNS_IMG, PURPLE, WHITE, GREY_DARK,
    MARGIN_MM, CONTENT_W_MM,
    set_run_font, no_borders, set_col_width, set_table_width,
)

# A4 portrait in twips
A4_W_TW = int(210 / 25.4 * 1440)
A4_H_TW = int(297 / 25.4 * 1440)
MARGIN_TW = int(MARGIN_MM / 25.4 * 1440)


# ── Image preparation ─────────────────────────────────────────────────────────

def _prepare_images(town_slug: str, cover_file: str):
    """Ensure composite images exist, generating from source files if needed."""
    try:
        from PIL import Image
        import numpy as np
    except ImportError:
        print("    WARNING: Pillow not installed. Run: pip install Pillow")
        return None, None

    cover_path   = TOWNS_IMG / town_slug / cover_file
    flourish_src = SHARED_IMG / "UQ_Flourish_Placeholder.png"
    logo_src     = SHARED_IMG / "UQlogo-Purple-cmyk.jpg"
    logo_out     = SHARED_IMG / "uq_logo.png"
    flourish_out = SHARED_IMG / "flourish_photo_masked.png"

    # Convert CMYK logo to RGB PNG
    if not logo_out.exists():
        if logo_src.exists():
            Image.open(logo_src).convert("RGB").save(str(logo_out))
            print(f"    Generated: {logo_out.name}")
        else:
            print(f"    WARNING: Logo not found at {logo_src}")
            logo_out = None

    # Composite: photo masked by flourish wave shape
    if not flourish_out.exists():
        if flourish_src.exists() and cover_path.exists():
            mask  = Image.open(flourish_src).convert("L")
            photo = Image.open(cover_path).convert("RGB")
            fw, fh = mask.size

            # Fit photo to flourish dimensions, crop centre
            pa = photo.width / photo.height
            ta = fw / fh
            if pa > ta:
                new_h = fh
                new_w = int(pa * fh)
                r = photo.resize((new_w, new_h), Image.LANCZOS)
                x = (new_w - fw) // 2
                r = r.crop((x, 0, x + fw, fh))
            else:
                new_w = fw
                new_h = int(fw / pa)
                r = photo.resize((new_w, new_h), Image.LANCZOS)
                y = (new_h - fh) // 3
                r = r.crop((0, y, fw, y + fh))

            import numpy as np
            arr   = np.array(mask)
            alpha = (arr < 50).astype("uint8") * 255
            rgba  = r.convert("RGBA")
            data  = np.array(rgba)
            data[:, :, 3] = alpha
            result = Image.fromarray(data)
            # Downscale to 1800px wide
            if result.width > 1800:
                scale  = 1800 / result.width
                result = result.resize(
                    (1800, int(result.height * scale)), Image.LANCZOS)
            result.save(str(flourish_out))
            print(f"    Generated: {flourish_out.name} ({result.size[0]}x{result.size[1]}px)")
        else:
            if not flourish_src.exists():
                print(f"    WARNING: Flourish mask not found: {flourish_src}")
            if not cover_path.exists():
                print(f"    WARNING: Cover photo not found: {cover_path}")
            flourish_out = None

    logo_out_p     = Path(logo_out)     if logo_out     else None
    flourish_out_p = Path(flourish_out) if flourish_out else None
    return (
        logo_out_p     if logo_out_p     and logo_out_p.exists()     else None,
        flourish_out_p if flourish_out_p and flourish_out_p.exists() else None,
    )


# ── Helpers ───────────────────────────────────────────────────────────────────

def _p(doc, space_before=0, space_after=0, line_pts=14, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a paragraph with exact line spacing to prevent Word inflation."""
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing      = Pt(line_pts)
    p.alignment          = align
    return p


def _gap(doc, pts):
    """Fixed-height blank line."""
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing      = Pt(pts)


# ── Page builder ──────────────────────────────────────────────────────────────

def build_cover_page(doc: Document, town_cfg: dict, date_str: str = "April 2026"):
    town_name  = town_cfg["name"]
    town_slug  = town_cfg["slug"]
    cover_file = town_cfg.get("cover_photo", "cover.jpg")

    logo_path, flourish_path = _prepare_images(town_slug, cover_file)

    # ── TOP BAR: 2-col table ──────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    set_table_width(tbl, Inches(CONTENT_W_MM / 25.4))

    lc = tbl.cell(0, 0)
    rc = tbl.cell(0, 1)

    lc.paragraphs[0].clear()
    lp1 = lc.paragraphs[0]
    lp1.paragraph_format.space_before = Pt(0)
    lp1.paragraph_format.space_after  = Pt(1)
    lp1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    lp1.paragraph_format.line_spacing      = Pt(11)
    set_run_font(lp1.add_run("Centre for Natural Gas"), 9, bold=True, colour=PURPLE)

    lp2 = lc.add_paragraph()
    lp2.paragraph_format.space_before = Pt(0)
    lp2.paragraph_format.space_after  = Pt(0)
    lp2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    lp2.paragraph_format.line_spacing      = Pt(11)
    set_run_font(lp2.add_run(date_str), 9, bold=True, colour=PURPLE)

    rc.paragraphs[0].clear()
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(0)
    rp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    rp.paragraph_format.line_spacing      = Pt(20)
    if logo_path:
        rp.add_run().add_picture(str(logo_path), width=Inches(1.4))

    content_tw = int(CONTENT_W_MM / 25.4 * 1440)
    left_w     = int(content_tw * 0.60)
    right_w    = content_tw - left_w
    set_col_width(lc, left_w)
    set_col_width(rc, right_w)
    set_table_width(tbl, Inches(CONTENT_W_MM / 25.4))
    no_borders(lc)
    no_borders(rc)

    # ── TITLE ─────────────────────────────────────────────────────────────────
    _gap(doc, 16)

    for line in [
        "Research Project: Cumulative",
        "social and economic impacts of",
        f"CSG development in {town_name}",
    ]:
        tp = _p(doc, space_before=0, space_after=2, line_pts=28)
        set_run_font(tp.add_run(line), 22, bold=True, colour=PURPLE)

    # ── COVER IMAGE ───────────────────────────────────────────────────────────
    _gap(doc, 10)

    ip = _p(doc, space_before=0, space_after=0, line_pts=14)
    if flourish_path:
        ip.add_run().add_picture(
            str(flourish_path), width=Inches(CONTENT_W_MM / 25.4))
    else:
        cover_path = TOWNS_IMG / town_slug / cover_file
        if cover_path.exists():
            ip.add_run().add_picture(
                str(cover_path), width=Inches(CONTENT_W_MM / 25.4))
        else:
            set_run_font(
                ip.add_run(
                    f"[Cover image missing — add {cover_file} to "
                    f"resources/images/towns/{town_slug}/]"),
                9, italic=True, colour=GREY_DARK)


# ── Standalone runner ─────────────────────────────────────────────────────────

def _load_town(town_name: str) -> dict:
    toml_path = _REPO_ROOT / "towns.toml"
    with open(toml_path, "rb") as f:
        cfg = tomllib.load(f)
    for slug, town in cfg.get("towns", {}).items():
        if town.get("name", "").lower() == town_name.lower() or slug.lower() == town_name.lower():
            return {
                "slug":        slug,
                "name":        town.get("name", town_name),
                "cover_photo": town.get("cover_photo", "cover.jpg"),
            }
    raise ValueError(f"Town '{town_name}' not found in towns.toml")


def _set_portrait(doc):
    """Set document to A4 portrait."""
    for section in doc.sections:
        section.page_width    = int(210 / 25.4 * 914400)
        section.page_height   = int(297 / 25.4 * 914400)
        section.left_margin   = Cm(MARGIN_MM / 10)
        section.right_margin  = Cm(MARGIN_MM / 10)
        section.top_margin    = Cm(MARGIN_MM / 10)
        section.bottom_margin = Cm(MARGIN_MM / 10)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate cover page standalone")
    parser.add_argument("--town", default="Chinchilla")
    parser.add_argument("--date", default=None)
    args = parser.parse_args()

    date_str = args.date or datetime.now().strftime("%B %Y")
    town_cfg = _load_town(args.town)

    print(f"Building cover page for: {town_cfg['name']}")

    doc = Document()
    _set_portrait(doc)
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    build_cover_page(doc, town_cfg, date_str)

    out_dir  = _REPO_ROOT / "booklets" / town_cfg["slug"]
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{town_cfg['name']}_cover_{datetime.now().year}.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")