"""
transform/booklet/make_booklet.py
----------------------------------
Entry point for booklet generation.

Page orientation:
    Pages 1-2  (cover, title)   : Portrait  A4
    Pages 3-10 (map, data pages): Landscape A4
    Final page (back)           : Portrait  A4  [TODO]

OOXML section model used here:
    A pPr-level sectPr closes the section ENDING at that paragraph.
    The body-level sectPr closes the final section.
    So we:
      1. Build portrait content (pages 1-2)
      2. Insert a pPr sectPr (portrait) to close section 1
      3. Build landscape content (pages 3-10)
      4. The body sectPr (landscape) closes section 2

Usage:
    python transform/booklet/make_booklet.py --town Chinchilla
    python transform/booklet/make_booklet.py --town Roma --date "June 2026"

Output:
    booklets/{slug}/{TownName}_Indicators_{year}.docx
"""
from __future__ import annotations

import argparse
import sys
import tomllib
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT  = SCRIPT_DIR.parent.parent

if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from common import BOOKLETS, MARGIN_MM
from pages.cover      import build_cover_page
from pages.title      import build_title_page
from pages.map        import build_map_page
from pages.population import build_population_page
from pages.data_page  import (build_unemployment_page, build_house_price_page,
                               build_rent_page, build_approvals_page,
                               build_rainfall_page, build_crime_page)

# A4 dimensions in twips (1440 twips = 1 inch = 25.4mm)
A4_SHORT_TW = int(210 / 25.4 * 1440)   # 11905  portrait width / landscape height
A4_LONG_TW  = int(297 / 25.4 * 1440)   # 16837  portrait height / landscape width
MARGIN_TW   = int(MARGIN_MM / 25.4 * 1440)


# ── Town config ───────────────────────────────────────────────────────────────

def load_town_cfg(town_name: str) -> dict:
    toml_path = REPO_ROOT / "towns.toml"
    with open(toml_path, "rb") as f:
        cfg = tomllib.load(f)
    for slug, town in cfg.get("towns", {}).items():
        if town.get("name", "").lower() == town_name.lower() or slug.lower() == town_name.lower():
            return {
                "slug":         slug,
                "name":         town.get("name", town_name),
                "sa2_code":     str(town.get("sa2_code", "")),
                "sa2_name":     town.get("sa2_name", town.get("name", town_name)),
                "district":     town.get("district", f"{town.get('name', town_name)} and District"),
                "cover_photo":  town.get("cover_photo", "cover.jpg"),
                "bom_station":  str(town.get("bom_station", "")),
                "qps_division": town.get("qps_division", ""),
                "sa2_map_url":  town.get("sa2_map_url", ""),
            }
    raise ValueError(f"Town '{town_name}' not found in towns.toml")


# ── Section helpers ───────────────────────────────────────────────────────────
#
# python-docx merges/drops the body sectPr when a pPr sectPr is present.
# We work around this by:
#   1. Inserting a pPr sectPr (portrait) to close section 1 (pages 1-2)
#   2. Post-processing the saved ZIP to inject the landscape body sectPr
#      directly into document.xml before </w:body>
#
# This two-step approach is necessary because python-docx's CT_Body
# serializer absorbs the body sectPr into the pPr sectPr during save().

def _insert_portrait_break(doc: Document) -> None:
    """Insert a pPr sectPr (portrait A4) to close the portrait section."""
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pPr    = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    pgSz   = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), str(A4_SHORT_TW))
    pgSz.set(qn("w:h"), str(A4_LONG_TW))
    # no orient attr = portrait
    sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    m = str(MARGIN_TW)
    for attr, val in [("w:top",m),("w:right",m),("w:bottom",m),("w:left",m),
                      ("w:header","720"),("w:footer","720"),("w:gutter","0")]:
        pgMar.set(qn(attr), val)
    sectPr.append(pgMar)
    pPr.append(sectPr)


def _fix_body_section(docx_path: Path) -> None:
    """
    Post-process the saved docx ZIP to inject a landscape body sectPr.
    python-docx drops the body sectPr when a pPr sectPr is present,
    so we inject it directly into the document.xml XML string.
    """
    import zipfile, shutil
    tmp = docx_path.with_suffix(".tmp.docx")
    shutil.copy(docx_path, tmp)

    with zipfile.ZipFile(tmp, "r") as zin:
        files = {n: zin.read(n) for n in zin.namelist()}

    xml = files["word/document.xml"].decode("utf-8")
    m   = str(MARGIN_TW)
    landscape_sect = (
        f'<w:sectPr>'
        f'<w:pgSz w:w="{A4_LONG_TW}" w:h="{A4_SHORT_TW}" w:orient="landscape"/>'
        f'<w:pgMar w:top="{m}" w:right="{m}" w:bottom="{m}" w:left="{m}"'
        f' w:header="720" w:footer="720" w:gutter="0"/>'
        f'</w:sectPr>'
    )
    xml = xml.replace("</w:body>", landscape_sect + "</w:body>")
    files["word/document.xml"] = xml.encode("utf-8")

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    tmp.unlink()


def _plain_page_break(doc: Document):
    """Simple page break within the current section."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── Document bootstrap ────────────────────────────────────────────────────────

def new_document() -> Document:
    """
    Create a blank document. We do NOT set the initial section here —
    it will be set by _insert_section_break (portrait) and
    _set_body_section (landscape) at assembly time.
    """
    doc = Document()
    # Remove python-docx's default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    return doc


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate town indicator booklet")
    parser.add_argument("--town",  required=True)
    parser.add_argument("--date",  default=None)
    parser.add_argument("--pages", nargs="+", default=["all"])
    args = parser.parse_args()

    date_str  = args.date or datetime.now().strftime("%B %Y")
    town_cfg  = load_town_cfg(args.town)
    slug      = town_cfg["slug"]
    pages     = args.pages
    build_all = "all" in pages

    print(f"Generating booklet for: {town_cfg['name']} ({slug})")
    print(f"Date: {date_str}")

    doc = new_document()

    # ── PORTRAIT content: pages 1 & 2 ────────────────────────────────────────
    if build_all or "cover" in pages:
        print("  [portrait]  Page 1: cover...")
        build_cover_page(doc, town_cfg, date_str)

    if build_all or "title" in pages:
        _plain_page_break(doc)
        print("  [portrait]  Page 2: title...")
        build_title_page(doc, town_cfg, date_str)

    # Close portrait section — pages 1-2 are portrait A4
    _insert_portrait_break(doc)

    # ── LANDSCAPE content: pages 3-10 ────────────────────────────────────────
    if build_all or "map" in pages:
        print("  [landscape] Page 3: map...")
        build_map_page(doc, town_cfg)

    if build_all or "population" in pages:
        _plain_page_break(doc)
        print("  [landscape] Page 4: population...")
        build_population_page(doc, town_cfg)

    if build_all or "unemployment" in pages:
        _plain_page_break(doc)
        print("  [landscape] Page 5: unemployment...")
        build_unemployment_page(doc, town_cfg)

    if build_all or "housing" in pages:
        _plain_page_break(doc)
        print("  [landscape] Page 6: housing...")
        build_house_price_page(doc, town_cfg)

    if build_all or "rent" in pages:
        _plain_page_break(doc)
        print("  [landscape] Page 7: rent...")
        build_rent_page(doc, town_cfg)

    if build_all or "approvals" in pages:
        _plain_page_break(doc)
        print("  [landscape] Page 8: approvals...")
        build_approvals_page(doc, town_cfg)

    if build_all or "rainfall" in pages:
        _plain_page_break(doc)
        print("  [landscape] Page 9: rainfall...")
        build_rainfall_page(doc, town_cfg)

    if build_all or "crime" in pages:
        _plain_page_break(doc)
        print("  [landscape] Page 10: crime...")
        build_crime_page(doc, town_cfg)

    # ── Save ──────────────────────────────────────────────────────────────────
    year     = datetime.now().year
    out_dir  = BOOKLETS / slug
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{town_cfg['name']}_Indicators_{year}.docx"
    doc.save(str(out_path))
    # Post-process: inject landscape body sectPr into saved ZIP
    # (python-docx drops body sectPr when pPr sectPr is present)
    _fix_body_section(out_path)
    print(f"\nSaved: {out_path}")


if __name__ == "__main__":
    main()