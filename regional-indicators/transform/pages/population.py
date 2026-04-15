"""
transform/booklet/pages/population.py
--------------------------------------
Page 4: Population chart.

Renders a bar chart of UCL resident population using python-docx table bars,
matching the visual style of the 2022 booklet.

Data source: cache/population/{slug}_population_ucl.json
  key: population_by_year → {year: count}
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

import sys
_BOOKLET_DIR = Path(__file__).resolve().parent.parent
if str(_BOOKLET_DIR) not in sys.path:
    sys.path.insert(0, str(_BOOKLET_DIR))
from common import (
    OUTPUT_DIR, REPO_ROOT, PURPLE, TEAL, WHITE, GREY_DARK, GREY_LIGHT,
    HEX_PURPLE, HEX_TEAL, HEX_WHITE, HEX_GREY_LIGHT, HEX_GREY_MED,
    CONTENT_W_MM,
    add_spacer, set_run_font, add_section_header, add_footer_line,
    no_borders, set_cell_bg, set_col_width, set_table_width,
)

# ── Layout constants ───────────────────────────────────────────────────────────
# All widths in twips (1440 twips = 1 inch = 25.4mm)
CONTENT_TWIPS = int(CONTENT_W_MM / 25.4 * 1440)  # ~9637 twips for 170mm
YEAR_COL_W    = 480    # year label column
VAL_COL_W     = 700    # value label column
BAR_AREA_W    = CONTENT_TWIPS - YEAR_COL_W - VAL_COL_W
ROW_HEIGHT_PT = 14     # bar row height in points (exact)

# ── Data loader ────────────────────────────────────────────────────────────────

def load_population(slug: str) -> dict[str, int]:
    cache = REPO_ROOT / "cache" / "population" / f"{slug}_population_ucl.json"
    if not cache.exists():
        return {}
    with open(cache) as f:
        data = json.load(f)
    return {yr: int(v) for yr, v in data.get("population_by_year", {}).items() if v}


# ── Chart builder ──────────────────────────────────────────────────────────────

def _bar_row(table, year: str, value: int, max_val: int,
             bar_hex: str, shade_row: bool):
    """Add one data row to the chart table."""
    row = table.add_row()

    # Scale bar width to available space
    bar_w   = max(1, int((value / max_val) * BAR_AREA_W))
    empty_w = BAR_AREA_W - bar_w

    # ── Year cell ──────────────────────────────────────────────────────────────
    yr_cell = row.cells[0]
    no_borders(yr_cell)
    if shade_row:
        set_cell_bg(yr_cell, HEX_GREY_LIGHT)
    yr_cell.paragraphs[0].clear()
    yp = yr_cell.paragraphs[0]
    yp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    yp.paragraph_format.space_before = Pt(0)
    yp.paragraph_format.space_after  = Pt(0)
    yr = yp.add_run(year)
    set_run_font(yr, 7, colour=GREY_DARK)

    # ── Bar cell (contains a nested 1-row table for the coloured bar) ──────────
    bar_cell = row.cells[1]
    no_borders(bar_cell)
    bar_cell.paragraphs[0].clear()

    # Build inner table: [coloured bar | empty space]
    col_widths = [bar_w, empty_w] if empty_w > 0 else [bar_w]
    inner = bar_cell.add_table(rows=1, cols=len(col_widths))
    _set_tbl_width(inner, BAR_AREA_W)

    filled_cell = inner.cell(0, 0)
    _set_tc_width(filled_cell, bar_w)
    set_cell_bg(filled_cell, bar_hex)
    no_borders(filled_cell)
    filled_cell.paragraphs[0].clear()
    fp = filled_cell.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    fp.add_run("")

    if empty_w > 0:
        empty_cell = inner.cell(0, 1)
        _set_tc_width(empty_cell, empty_w)
        no_borders(empty_cell)
        if shade_row:
            set_cell_bg(empty_cell, HEX_GREY_LIGHT)
        empty_cell.paragraphs[0].clear()
        ep = empty_cell.paragraphs[0]
        ep.paragraph_format.space_before = Pt(0)
        ep.paragraph_format.space_after  = Pt(0)
        ep.add_run("")

    # Set inner row height
    _set_row_height(inner.rows[0], ROW_HEIGHT_PT)

    # ── Value cell ─────────────────────────────────────────────────────────────
    val_cell = row.cells[2]
    no_borders(val_cell)
    if shade_row:
        set_cell_bg(val_cell, HEX_GREY_LIGHT)
    val_cell.paragraphs[0].clear()
    vp = val_cell.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    vp.paragraph_format.space_before = Pt(0)
    vp.paragraph_format.space_after  = Pt(0)
    vr = vp.add_run(f"{value:,}")
    set_run_font(vr, 7, colour=GREY_DARK)

    # Set outer row height
    _set_row_height(row, ROW_HEIGHT_PT)


def _set_row_height(row, height_pt):
    tr    = row._tr
    trPr  = tr.get_or_add_trPr()
    for existing in trPr.findall(qn("w:trHeight")):
        trPr.remove(existing)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(height_pt * 20)))  # twips = pt * 20
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def _set_tbl_width(table, width_twips):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for ex in tblPr.findall(qn("w:tblW")):
        tblPr.remove(ex)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(int(width_twips)))
    tblW.set(qn("w:type"), "dxa")
    # tblW must come after tblStyle but before jc/tblBorders in schema
    # Insert after tblStyle if present, else at position 0
    tbl_style = tblPr.find(qn("w:tblStyle"))
    if tbl_style is not None:
        tbl_style.addnext(tblW)
    else:
        tblPr.insert(0, tblW)


def _set_tc_width(cell, width_twips):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:tcW")):
        tcPr.remove(ex)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_twips)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.insert(0, tcW)  # tcW must be first in tcPr


def _chart_header_row(table, y_label: str):
    """Add a column header row above the data bars."""
    hdr = table.add_row()

    yr_cell = hdr.cells[0]
    no_borders(yr_cell)
    set_cell_bg(yr_cell, HEX_PURPLE)
    yr_cell.paragraphs[0].clear()
    yp = yr_cell.paragraphs[0]
    yp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    yp.paragraph_format.space_before = Pt(1)
    yp.paragraph_format.space_after  = Pt(1)
    set_run_font(yp.add_run("Year"), 7, bold=True, colour=WHITE)

    bar_cell = hdr.cells[1]
    no_borders(bar_cell)
    set_cell_bg(bar_cell, HEX_PURPLE)
    bar_cell.paragraphs[0].clear()
    bp = bar_cell.paragraphs[0]
    bp.paragraph_format.space_before = Pt(1)
    bp.paragraph_format.space_after  = Pt(1)
    set_run_font(bp.add_run(y_label), 7, bold=True, colour=WHITE)

    val_cell = hdr.cells[2]
    no_borders(val_cell)
    set_cell_bg(val_cell, HEX_PURPLE)
    val_cell.paragraphs[0].clear()
    vp = val_cell.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    vp.paragraph_format.space_before = Pt(1)
    vp.paragraph_format.space_after  = Pt(1)
    set_run_font(vp.add_run("Count"), 7, bold=True, colour=WHITE)


def build_chart_table(doc, series: dict[str, int], y_label: str, bar_hex: str):
    """
    Build and add a bar chart table to `doc`.
    series: {year_str: numeric_value}
    """
    if not series:
        p = doc.add_paragraph()
        set_run_font(p.add_run("⚠  No data available for this indicator."), 9,
                     italic=True, colour=GREY_DARK)
        return

    years   = sorted(series.keys())
    max_val = max(series.values()) * 1.05   # 5% headroom

    # Outer 3-column table: [year | bar area | value]
    table = doc.add_table(rows=0, cols=3)
    _set_tbl_width(table, CONTENT_TWIPS)
    _set_tc_width(table.columns[0].cells[0] if table.rows else table.add_row().cells[0],
                  YEAR_COL_W)

    # Header
    _chart_header_row(table, y_label)

    # Set fixed column widths on header row
    hdr_row = table.rows[0]
    _set_tc_width(hdr_row.cells[0], YEAR_COL_W)
    _set_tc_width(hdr_row.cells[1], BAR_AREA_W)
    _set_tc_width(hdr_row.cells[2], VAL_COL_W)
    _set_row_height(hdr_row, ROW_HEIGHT_PT + 2)

    # Data rows
    for i, year in enumerate(years):
        _bar_row(table, year, int(series[year]), max_val, bar_hex, shade_row=(i % 2 == 1))

    return table


# ── Page builder ──────────────────────────────────────────────────────────────

def build_population_page(doc: Document, town_cfg: dict):
    """
    Append population chart page to `doc`.
    """
    town_name = town_cfg["name"]
    slug      = town_cfg["slug"]

    add_section_header(doc, "Population",
                       f"Estimated resident population — {town_name} Urban Centre and Locality (UCL)")
    add_spacer(doc, 10)

    pop_data = load_population(slug)

    if pop_data:
        build_chart_table(doc, pop_data,
                          y_label="No. of persons (UCL residents)",
                          bar_hex=HEX_PURPLE)
    else:
        p = doc.add_paragraph()
        set_run_font(
            p.add_run(f"Population data not available for {town_name}. "
                      f"Ensure cache/population/{slug}_population_ucl.json exists."),
            9, italic=True, colour=GREY_DARK
        )

    add_spacer(doc, 8)

    # ── Data notes ────────────────────────────────────────────────────────────
    notes = [
        ("Source: ", True),
        ("QGSO Estimated Resident Population by Urban Centre and Locality (UCL), Queensland.", False),
        ("  Boundary: ", True),
        (f"Urban centre and locality boundary ({town_name} UCL).", False),
        ("  Data: ", True),
        ("Annual ERP at 30 June each year.", False),
    ]
    dn_p = doc.add_paragraph()
    dn_p.paragraph_format.space_before = Pt(4)
    dn_p.paragraph_format.space_after  = Pt(2)
    for text, bold in notes:
        r = dn_p.add_run(text)
        set_run_font(r, 8, bold=bold, colour=GREY_DARK)

    add_footer_line(doc, town_name)