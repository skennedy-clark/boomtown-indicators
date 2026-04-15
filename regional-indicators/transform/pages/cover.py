"""
transform/booklet/pages/cover.py
---------------------------------
Page 1: Cover page.

Layout (top → bottom):
  - Top bar: "Centre for Natural Gas" + date (left), UQ logo (right)
  - Large title text in UQ purple
  - Cover photo masked by the UQ flourish wave shape (fills lower portion)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
_BOOKLET_DIR = Path(__file__).resolve().parent.parent
if str(_BOOKLET_DIR) not in sys.path:
    sys.path.insert(0, str(_BOOKLET_DIR))
from common import (
    SHARED_IMG, TOWNS_IMG, PURPLE, WHITE, GREY_DARK,
    MARGIN_MM, CONTENT_W_MM,
    add_spacer, set_run_font, no_borders,
    set_col_width, set_table_width,
)


def build_cover_page(doc: Document, town_cfg: dict, date_str: str = "April 2026"):
    """
    Append cover page content to `doc`.

    Args:
        doc:       python-docx Document (page margins already set by caller)
        town_cfg:  dict with keys: name, cover_photo (filename in towns/{slug}/)
        date_str:  date string shown on cover, e.g. "April 2026"
    """
    town_name  = town_cfg["name"]
    town_slug  = town_cfg["slug"]
    cover_file = town_cfg.get("cover_photo", "cover.jpg")

    cover_path    = TOWNS_IMG / town_slug / cover_file
    logo_path     = SHARED_IMG / "uq_logo.png"
    flourish_path = SHARED_IMG / "flourish_photo_masked.png"

    # ── TOP BAR: two-column table (Centre info left | UQ logo right) ───────────
    # Use a 2-col borderless table so logo can sit right-aligned
    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT = 1  # center
    set_table_width(top_table, Inches(CONTENT_W_MM / 25.4))

    left_cell  = top_table.cell(0, 0)
    right_cell = top_table.cell(0, 1)
    # no_borders called AFTER set_col_width so tcBorders follows tcW in tcPr
    # (will be called below after column widths are set)

    # Left: Centre name + date
    left_cell.paragraphs[0].clear()
    lp = left_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(2)
    r = lp.add_run("Centre for Natural Gas")
    set_run_font(r, 10, bold=True, colour=PURPLE)

    lp2 = left_cell.add_paragraph()
    lp2.paragraph_format.space_before = Pt(0)
    lp2.paragraph_format.space_after  = Pt(0)
    r2 = lp2.add_run(date_str)
    set_run_font(r2, 10, bold=True, colour=PURPLE)

    # Right: UQ logo, right-aligned
    right_cell.paragraphs[0].clear()
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(0)
    rrun = rp.add_run()
    if logo_path.exists():
        rrun.add_picture(str(logo_path), width=Inches(1.6))

    # Column widths: left 60%, right 40%
    content_twips = int(CONTENT_W_MM / 25.4 * 1440)
    left_w  = int(content_twips * 0.60)
    right_w = content_twips - left_w
    set_col_width(left_cell,  left_w)
    set_col_width(right_cell, right_w)
    set_table_width(top_table, Inches(CONTENT_W_MM / 25.4))
    # Apply borders after col widths so tcBorders comes after tcW in tcPr
    no_borders(left_cell)
    no_borders(right_cell)

    # ── TITLE TEXT ─────────────────────────────────────────────────────────────
    add_spacer(doc, 24)

    title_lines = [
        f"Research Project: Cumulative",
        f"social and economic impacts of",
        f"CSG development in {town_name}",
    ]
    for i, line in enumerate(title_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(line)
        set_run_font(run, 28, bold=True, colour=PURPLE)

    # ── COVER PHOTO (masked by flourish) ───────────────────────────────────────
    add_spacer(doc, 18)

    # The flourish-masked photo fills the content width
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    img_p.paragraph_format.space_before = Pt(0)
    img_p.paragraph_format.space_after  = Pt(0)
    img_run = img_p.add_run()

    if flourish_path.exists():
        img_run.add_picture(str(flourish_path), width=Inches(CONTENT_W_MM / 25.4))
    elif cover_path.exists():
        # Fallback: just use the raw photo if flourish composite not found
        img_run.add_picture(str(cover_path), width=Inches(CONTENT_W_MM / 25.4))