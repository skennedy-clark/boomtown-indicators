"""
transform/booklet/pages/map.py
-------------------------------
Page 3: SA2 boundary map page.

Automatically fetches the SA2 boundary map PDF from QGSO, rasterises it,
crops to the main map area, and embeds it in the booklet.

Map PDF URL pattern:
  https://www.qgso.qld.gov.au/geographies-maps/maps/qld-sa2-asgs-2021-{slug}.pdf
  where slug = sa2_name lowercased, spaces to hyphens, " - " to "-"

The rendered PNG is cached to:
  resources/images/towns/{slug}/sa2_map.png

so subsequent runs do not re-download. Delete the cached PNG to force refresh.

Page layout:
  - Map image (full content width)
  - Caption: "{sa2_name} statistical area 2 (SA2) boundary"
  - Source URL
  - Footer community line
"""
from __future__ import annotations

import re
import urllib.request
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
_BOOKLET_DIR = Path(__file__).resolve().parent.parent
if str(_BOOKLET_DIR) not in sys.path:
    sys.path.insert(0, str(_BOOKLET_DIR))
from common import (
    TOWNS_IMG, PURPLE, GREY_DARK, HEX_GREY_LIGHT, HEX_PURPLE,
    CONTENT_W_MM,
    add_spacer, set_run_font, add_footer_line,
    set_cell_bg, no_borders,
)

BASE_URL = "https://www.qgso.qld.gov.au/geographies-maps/maps/"
MAP_DPI  = 150


def _sa2_url_slug(sa2_name: str) -> str:
    slug = sa2_name.lower()
    slug = slug.replace(" - ", "-").replace(" - ", "-")
    slug = slug.replace(" ", "-")
    slug = slug.replace("'", "").replace("(", "").replace(")", "")
    slug = re.sub(r"-+", "-", slug).strip("-")
    return slug


def _map_url(sa2_name: str) -> str:
    return BASE_URL + f"qld-sa2-asgs-2021-{_sa2_url_slug(sa2_name)}.pdf"


def _fetch_and_render(sa2_name: str, out_png: Path) -> bool:
    try:
        import pypdfium2 as pdfium
        from PIL import Image
    except ImportError:
        print("    WARNING: pypdfium2 or Pillow not installed.")
        print("    Run: pip install pypdfium2 Pillow")
        return False

    url     = _map_url(sa2_name)
    pdf_tmp = out_png.parent / "_sa2_map_tmp.pdf"
    out_png.parent.mkdir(parents=True, exist_ok=True)

    print(f"    Fetching SA2 map: {url}")
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            pdf_tmp.write_bytes(resp.read())
    except Exception as e:
        print(f"    WARNING: Could not fetch map PDF: {e}")
        print(f"    URL: {url}")
        print(f"    Download manually and save as: {out_png}")
        pdf_tmp.unlink(missing_ok=True)
        return False

    try:
        pdf    = pdfium.PdfDocument(str(pdf_tmp))
        page   = pdf[0]
        bitmap = page.render(scale=MAP_DPI / 72)
        img    = bitmap.to_pil().convert("RGB")
        w, h   = img.size
        # QGSO PDF layout (A4 landscape):
        #   top ~8%  : title bar (stripped — heading is in docx)
        #   left ~65%: main map area
        #   right 35%: QLD state inset + legend + credits (excluded)
        title_h  = int(h * 0.08)
        map_crop = img.crop((0, title_h, int(w * 0.67), h))
        map_crop.save(str(out_png))
        pdf_tmp.unlink(missing_ok=True)
        print(f"    Saved: {out_png.name} ({map_crop.size[0]}x{map_crop.size[1]}px)")
        return True
    except Exception as e:
        print(f"    WARNING: Could not rasterise map PDF: {e}")
        pdf_tmp.unlink(missing_ok=True)
        return False


def get_map_png(town_cfg: dict) -> Path | None:
    slug     = town_cfg["slug"]
    sa2_name = town_cfg.get("sa2_name", town_cfg["name"])
    out_png  = TOWNS_IMG / slug / "sa2_map.png"

    if out_png.exists():
        print(f"    Using cached map: {out_png}")
        return out_png

    success = _fetch_and_render(sa2_name, out_png)
    return out_png if success else None


def build_map_page(doc: Document, town_cfg: dict):
    town_name = town_cfg["name"]
    sa2_name  = town_cfg.get("sa2_name", town_name)
    sa2_code  = town_cfg.get("sa2_code", "")
    map_url   = _map_url(sa2_name)

    # Heading
    hd = doc.add_paragraph()
    hd.paragraph_format.space_before = Pt(0)
    hd.paragraph_format.space_after  = Pt(6)
    set_run_font(
        hd.add_run(
            f"Queensland Statistical Areas, Level 2 (SA2), 2021 — "
            f"{sa2_name} (ASGS Code {sa2_code})"
        ),
        11, bold=True, colour=PURPLE
    )

    # Map image or placeholder
    map_png = get_map_png(town_cfg)

    if map_png and map_png.exists():
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        img_p.paragraph_format.space_before = Pt(0)
        img_p.paragraph_format.space_after  = Pt(6)
        img_p.add_run().add_picture(str(map_png), width=Inches(CONTENT_W_MM / 25.4))
    else:
        _placeholder(doc, town_name, sa2_code, map_url)

    # Caption
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(2)
    set_run_font(
        cap.add_run(f"{sa2_name} statistical area 2 (SA2) boundary"),
        9, bold=True, colour=GREY_DARK
    )

    src = doc.add_paragraph()
    src.paragraph_format.space_before = Pt(0)
    src.paragraph_format.space_after  = Pt(0)
    set_run_font(src.add_run("Source:  "), 8, bold=True, colour=GREY_DARK)
    set_run_font(src.add_run(map_url),     8, colour=PURPLE)

    add_footer_line(doc, town_name)


def _placeholder(doc, town_name, sa2_code, map_url):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, HEX_GREY_LIGHT)

    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), HEX_PURPLE)
        tcBorders.append(el)
    tcW_el = tcPr.find(qn("w:tcW"))
    if tcW_el is not None:
        tcW_el.addnext(tcBorders)
    else:
        tcPr.append(tcBorders)

    tr   = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(110 / 25.4 * 1440)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)

    items = [
        (f"[ SA2 MAP — {town_name.upper()} ]",                            True,  12, 36),
        (f"SA2 Code: {sa2_code}",                                          False,  9,  6),
        ("Map could not be fetched automatically.",                         True,   8, 10),
        ("Download from:",                                                  False,  8,  4),
        (map_url,                                                           False,  8,  2),
        (f"Save as: resources/images/towns/{town_name.lower()}/sa2_map.png", False, 8,  2),
        (f"Then re-run: python transform/booklet/make_booklet.py "
         f"--town {town_name}",                                             False,  8,  2),
    ]
    first = True
    for text, bold, size, top_space in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(top_space)
        p.paragraph_format.space_after  = Pt(0)
        colour = PURPLE if bold else GREY_DARK
        set_run_font(p.add_run(text), size, bold=bold, colour=colour)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)