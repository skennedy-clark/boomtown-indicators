"""
transform/booklet/pages/title.py
---------------------------------
Page 2: Title / contact details page.

Layout:
  - UQ logo top right
  - "DATA REPORT" label
  - Large italic title "INDICATORS OF CHANGE IN {TOWN} AND DISTRICT"
  - Centred date line
  - Contact block at bottom
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
_BOOKLET_DIR = Path(__file__).resolve().parent.parent
if str(_BOOKLET_DIR) not in sys.path:
    sys.path.insert(0, str(_BOOKLET_DIR))
from common import (
    SHARED_IMG, PURPLE, WHITE, GREY_DARK,
    CONTENT_W_MM,
    add_para, add_spacer, set_run_font, no_borders,
    set_table_width,
)

# Contact details — update as needed
CONTACT = {
    "name":       "UQ Centre for Natural Gas",
    "university": "The University of Queensland | Brisbane, QLD, 4072 AUSTRALIA",
    "web":        "https://natural-gas.centre.uq.edu.au/",
    "email":      "natural.gas@uq.edu.au",
}


def build_title_page(doc: Document, town_cfg: dict, date_str: str = "April 2026"):
    """
    Append title/contact page content to `doc`.

    Args:
        doc:       python-docx Document
        town_cfg:  dict with keys: name, district (optional, defaults to "{name} and District")
        date_str:  update date string
    """
    town_name  = town_cfg["name"]
    district   = town_cfg.get("district", f"{town_name} and District")
    logo_path  = SHARED_IMG / "uq_logo.png"

    # ── UQ LOGO (top right) ────────────────────────────────────────────────────
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.space_after  = Pt(4)
    logo_run = logo_p.add_run()
    if logo_path.exists():
        logo_run.add_picture(str(logo_path), width=Inches(1.8))

    # "CREATE CHANGE" tagline under logo, right-aligned
    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tagline_p.paragraph_format.space_before = Pt(0)
    tagline_p.paragraph_format.space_after  = Pt(0)
    tr = tagline_p.add_run("CREATE CHANGE")
    set_run_font(tr, 7, colour=GREY_DARK)
    # Thin rule under tagline
    pPr = tagline_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:color"), "AAAAAA")
    bot.set(qn("w:space"), "2")
    pBdr.append(bot)
    # pBdr must precede spacing/jc in pPr — insert at front
    pPr.insert(0, pBdr)

    # ── BIG SPACER ─────────────────────────────────────────────────────────────
    add_spacer(doc, 60)

    # ── "DATA REPORT" label ────────────────────────────────────────────────────
    dr_p = doc.add_paragraph()
    dr_p.paragraph_format.space_before = Pt(0)
    dr_p.paragraph_format.space_after  = Pt(24)
    dr_run = dr_p.add_run("DATA REPORT")
    set_run_font(dr_run, 14, bold=True, colour=PURPLE)

    # ── MAIN TITLE ─────────────────────────────────────────────────────────────
    title_lines = [
        "INDICATORS OF CHANGE IN",
        district.upper(),
    ]
    for line in title_lines:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(4)
        tr = tp.add_run(line)
        set_run_font(tr, 26, bold=True, italic=True, colour=PURPLE)

    # ── DATE LINE ──────────────────────────────────────────────────────────────
    add_spacer(doc, 48)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(0)
    date_p.paragraph_format.space_after  = Pt(0)
    dr2 = date_p.add_run(f"Statistical data updated in {date_str}")
    set_run_font(dr2, 10, colour=GREY_DARK)

    # ── CONTACT BLOCK ──────────────────────────────────────────────────────────
    add_spacer(doc, 48)

    c_label = doc.add_paragraph()
    c_label.paragraph_format.space_before = Pt(0)
    c_label.paragraph_format.space_after  = Pt(6)
    set_run_font(c_label.add_run("Contact:"), 10, colour=GREY_DARK)

    c_name = doc.add_paragraph()
    c_name.paragraph_format.space_before = Pt(0)
    c_name.paragraph_format.space_after  = Pt(2)
    set_run_font(c_name.add_run(CONTACT["name"]), 10, bold=True)

    c_uni = doc.add_paragraph()
    c_uni.paragraph_format.space_before = Pt(0)
    c_uni.paragraph_format.space_after  = Pt(2)
    set_run_font(c_uni.add_run(CONTACT["university"]), 10, colour=GREY_DARK)

    c_web = doc.add_paragraph()
    c_web.paragraph_format.space_before = Pt(0)
    c_web.paragraph_format.space_after  = Pt(2)
    wr = c_web.add_run(f"W: {CONTACT['web']}")
    set_run_font(wr, 10, colour=PURPLE)

    c_email = doc.add_paragraph()
    c_email.paragraph_format.space_before = Pt(0)
    c_email.paragraph_format.space_after  = Pt(0)
    er = c_email.add_run(f"E: {CONTACT['email']}")
    set_run_font(er, 10, colour=PURPLE)