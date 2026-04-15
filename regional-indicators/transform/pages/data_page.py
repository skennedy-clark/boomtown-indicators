"""
transform/booklet/pages/data_page.py
--------------------------------------
Generic chart page builder, used for all indicator pages:
  - Unemployment
  - Housing (sale price, sales volume, rent, approvals)
  - Rainfall
  - Crime (all offences, drug, good order, theft, traffic)

Each page follows the same pattern:
  - Purple section header banner
  - Bar chart table (one or two series)
  - Data notes paragraph
  - Footer line

The chart renders as a table of proportional bars, matching the visual
style of the 2022 PDF booklets.
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
_BOOKLET_DIR = Path(__file__).resolve().parent.parent
if str(_BOOKLET_DIR) not in sys.path:
    sys.path.insert(0, str(_BOOKLET_DIR))
from common import (
    PURPLE, TEAL, WHITE, GREY_DARK,
    HEX_PURPLE, HEX_TEAL, HEX_WHITE, HEX_GREY_LIGHT, HEX_GREY_MED,
    CONTENT_W_MM,
    add_spacer, set_run_font, add_section_header, add_footer_line,
    no_borders, set_cell_bg,
)

# ── Layout constants ───────────────────────────────────────────────────────────
CONTENT_TWIPS = int(CONTENT_W_MM / 25.4 * 1440)
YEAR_COL_W    = 480
VAL_COL_W     = 780
BAR_AREA_W    = CONTENT_TWIPS - YEAR_COL_W - VAL_COL_W
ROW_H_PT      = 14


# ── Low-level XML helpers (local, no EMU conversion needed) ──────────────────

def _tbl_width(table, twips: int):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for ex in tblPr.findall(qn("w:tblW")):
        tblPr.remove(ex)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(int(twips)))
    tblW.set(qn("w:type"), "dxa")
    tbl_style = tblPr.find(qn("w:tblStyle"))
    if tbl_style is not None:
        tbl_style.addnext(tblW)
    else:
        tblPr.insert(0, tblW)


def _tc_width(cell, twips: int):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:tcW")):
        tcPr.remove(ex)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(twips)))
    tcW.set(qn("w:type"), "dxa")
    # tcW must be first child of tcPr in schema order
    tcPr.insert(0, tcW)


def _row_height(row, pt: float):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    for ex in trPr.findall(qn("w:trHeight")):
        trPr.remove(ex)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(pt * 20)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def _cell_para(cell, text: str, size_pt: float, bold=False,
               colour=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size_pt, bold=bold, colour=colour or GREY_DARK)


# ── Chart header row ──────────────────────────────────────────────────────────

def _header_row(table, series_labels: list[str]):
    hdr = table.add_row()
    _row_height(hdr, ROW_H_PT + 3)

    # Year cell
    yc = hdr.cells[0]
    no_borders(yc)
    _tc_width(yc, YEAR_COL_W)
    set_cell_bg(yc, HEX_PURPLE)
    _cell_para(yc, "Year", 7, bold=True, colour=WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Bar label cell
    bc = hdr.cells[1]
    no_borders(bc)
    _tc_width(bc, BAR_AREA_W)
    set_cell_bg(bc, HEX_PURPLE)
    label = "  /  ".join(series_labels)
    _cell_para(bc, label, 7, bold=True, colour=WHITE)

    # Value cell
    vc = hdr.cells[2]
    no_borders(vc)
    _tc_width(vc, VAL_COL_W)
    set_cell_bg(vc, HEX_PURPLE)
    _cell_para(vc, "Value", 7, bold=True, colour=WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)


# ── Data row ──────────────────────────────────────────────────────────────────

def _data_row(table, year: str, values: list[float | None], max_val: float,
              bar_colours: list[str], format_fn, shade: bool,
              multi_series: bool = False):
    """
    Add one data row. Supports 1 or 2 series (stacked bars for 2-series).
    """
    row = table.add_row()
    _row_height(row, ROW_H_PT)
    row_bg = HEX_GREY_LIGHT if shade else HEX_WHITE

    # ── Year cell ─────────────────────────────────────────────────────────────
    yc = row.cells[0]
    no_borders(yc)
    _tc_width(yc, YEAR_COL_W)
    set_cell_bg(yc, row_bg)
    _cell_para(yc, year, 7, colour=GREY_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Bar cell ──────────────────────────────────────────────────────────────
    bc = row.cells[1]
    no_borders(bc)
    _tc_width(bc, BAR_AREA_W)
    bc.paragraphs[0].clear()
    bc.paragraphs[0].paragraph_format.space_before = Pt(0)
    bc.paragraphs[0].paragraph_format.space_after  = Pt(0)

    # Build inner bar table
    # Segments: one per series value, plus empty remainder
    segments = []
    total_bar = 0
    for val, colour in zip(values, bar_colours):
        if val is not None and val > 0:
            w = max(1, int((val / max_val) * BAR_AREA_W))
            segments.append((w, colour))
            total_bar += w
        else:
            segments.append((0, colour))

    empty_w = max(0, BAR_AREA_W - total_bar)
    all_segs = [(w, c) for w, c in segments if w > 0]
    if empty_w > 0:
        all_segs.append((empty_w, row_bg))

    if not all_segs:
        all_segs = [(BAR_AREA_W, row_bg)]

    col_widths = [w for w, _ in all_segs]
    inner = bc.add_table(rows=1, cols=len(col_widths))
    _tbl_width(inner, BAR_AREA_W)
    inner_row = inner.rows[0]
    _row_height(inner_row, ROW_H_PT)

    for ci, (w, colour) in enumerate(all_segs):
        cell = inner_row.cells[ci]
        _tc_width(cell, w)
        no_borders(cell)
        set_cell_bg(cell, colour)
        cell.paragraphs[0].clear()
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(0)
        cell.paragraphs[0].add_run("")

    # ── Value cell ────────────────────────────────────────────────────────────
    vc = row.cells[2]
    no_borders(vc)
    _tc_width(vc, VAL_COL_W)
    set_cell_bg(vc, row_bg)
    formatted = "  /  ".join(
        format_fn(v) if v is not None else "—"
        for v in values
    )
    _cell_para(vc, formatted, 7, colour=GREY_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)


# ── Public chart builder ──────────────────────────────────────────────────────

def build_chart(doc: Document,
                series: list[dict],
                format_fn,
                years: list[str] | None = None):
    """
    Build and add a bar chart table to `doc`.

    Args:
        doc:       python-docx Document
        series:    list of dicts, each with:
                     'label': str  — column header label
                     'data':  dict — {year_str: numeric_value}
                     'color': str  — hex colour for bar (without #)
        format_fn: callable(value) → str for value display
        years:     ordered list of year strings to include;
                   defaults to union of all series years, sorted
    """
    if not series:
        return

    # Collect all years across all series
    if years is None:
        all_years = set()
        for s in series:
            all_years |= set(s['data'].keys())
        years = sorted(all_years)

    # Max value across all series for bar scaling
    all_vals = [v for s in series for v in s['data'].values() if v is not None]
    if not all_vals:
        p = doc.add_paragraph()
        set_run_font(p.add_run("⚠  No data available."), 9, italic=True, colour=GREY_DARK)
        return
    max_val = max(all_vals) * 1.08  # 8% headroom

    multi = len(series) > 1
    labels = [s['label'] for s in series]
    colours = [s['color'] for s in series]

    # Outer 3-col table
    table = doc.add_table(rows=0, cols=3)
    _tbl_width(table, CONTENT_TWIPS)

    _header_row(table, labels)

    for i, year in enumerate(years):
        vals = [s['data'].get(year) for s in series]
        _data_row(table, year, vals, max_val, colours, format_fn,
                  shade=(i % 2 == 1), multi_series=multi)


# ── Format functions ──────────────────────────────────────────────────────────

def fmt_currency(v):
    return f"${int(round(v)):,}"

def fmt_currency_k(v):
    return f"${v/1000:.0f}k"

def fmt_int(v):
    return f"{int(round(v)):,}"

def fmt_pct(v):
    return f"{v:.1f}%"

def fmt_rate(v):
    return f"{v:.1f}"

def fmt_mm(v):
    return f"{v:.0f} mm"


# ── Page builders ─────────────────────────────────────────────────────────────

def build_unemployment_page(doc: Document, town_cfg: dict):
    import json
    from common import REPO_ROOT
    town_name = town_cfg["name"]
    slug      = town_cfg["slug"]

    add_section_header(doc, "Unemployment Rate",
                       f"Annual smoothed unemployment rate — {town_name} SA2")
    add_spacer(doc, 8)

    cache = REPO_ROOT / "cache" / "unemployment" / f"{slug}_salm.json"
    series_data = {}
    if cache.exists():
        with open(cache) as f:
            d = json.load(f)
        series_data = {yr: v for yr, v in
                       d.get("indicators", {}).get("unemployment", {}).items()
                       if v is not None}

    build_chart(doc,
        series=[{"label": "Unemployment rate (%)", "data": series_data, "color": HEX_TEAL}],
        format_fn=fmt_pct)

    add_spacer(doc, 8)
    _data_notes(doc, [
        ("Source: ", True),
        ("DEWR Small Area Labour Markets (SALM), smoothed SA2 quarterly estimates.", False),
        ("  Annual value: ", True),
        ("mean of 4 quarterly estimates.", False),
        ("  Note: ", True),
        ("SA2s below statistical threshold are excluded from SALM — "
         "data may start from 2010 or later.", False),
    ])
    add_footer_line(doc, town_name)


def build_house_price_page(doc: Document, town_cfg: dict):
    import json
    from common import REPO_ROOT
    town_name = town_cfg["name"]
    slug      = town_cfg["slug"]

    add_section_header(doc, "Housing — Median Sale Price & Volume",
                       f"Detached dwelling sales — {town_name} SA2")
    add_spacer(doc, 8)

    cache = REPO_ROOT / "cache" / "housing" / f"{slug}_qgso.json"
    price_data = {}
    sales_data = {}
    if cache.exists():
        with open(cache) as f:
            d = json.load(f)
        inds = d.get("indicators", {})
        price_data = {yr: v for yr, v in inds.get("housing_median_price", {}).items() if v}
        sales_data = {yr: v for yr, v in inds.get("housing_sales_count", {}).items() if v}

    # Sale price chart
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(4)
    set_run_font(sub_p.add_run("Median house sale price ($)"), 9, bold=True, colour=PURPLE)

    build_chart(doc,
        series=[{"label": "Median sale price", "data": price_data, "color": HEX_PURPLE}],
        format_fn=fmt_currency)

    add_spacer(doc, 10)

    # Sales volume chart
    sub_p2 = doc.add_paragraph()
    sub_p2.paragraph_format.space_before = Pt(0)
    sub_p2.paragraph_format.space_after  = Pt(4)
    set_run_font(sub_p2.add_run("Number of house sales"), 9, bold=True, colour=PURPLE)

    build_chart(doc,
        series=[{"label": "No. of sales", "data": sales_data, "color": HEX_TEAL}],
        format_fn=fmt_int)

    add_spacer(doc, 8)
    _data_notes(doc, [
        ("Source: ", True),
        ("QGSO Regional Database — Residential land and dwelling sales (DNRM).", False),
        ("  Boundary: ", True),
        (f"SA2 ({town_name}, ASGS 2021).", False),
        ("  Aggregation: ", True),
        ("Sales count = Dec-quarter rolling 12-month total. "
         "Median price = mean of 4 quarterly medians per calendar year.", False),
    ])
    add_footer_line(doc, town_name)


def build_rent_page(doc: Document, town_cfg: dict):
    import json
    from common import REPO_ROOT
    town_name = town_cfg["name"]
    slug      = town_cfg["slug"]

    add_section_header(doc, "Housing — Median Weekly Rent",
                       f"3-bedroom house median weekly rent — {town_name} SA2")
    add_spacer(doc, 8)

    cache = REPO_ROOT / "cache" / "housing" / f"{slug}_qgso.json"
    rent_data = {}
    if cache.exists():
        with open(cache) as f:
            d = json.load(f)
        rent_data = {yr: v for yr, v in
                     d.get("indicators", {}).get("rent_3bed_median", {}).items() if v}

    build_chart(doc,
        series=[{"label": "Median rent ($/week)", "data": rent_data, "color": HEX_PURPLE}],
        format_fn=fmt_currency)

    add_spacer(doc, 8)
    _data_notes(doc, [
        ("Source: ", True),
        ("QGSO Regional Database — Median rent (RTA).", False),
        ("  Boundary: ", True),
        (f"SA2 ({town_name}, ASGS 2021).", False),
        ("  Aggregation: ", True),
        ("Annual figure = mean of 4 quarterly medians.", False),
    ])
    add_footer_line(doc, town_name)


def build_approvals_page(doc: Document, town_cfg: dict):
    import json
    from common import REPO_ROOT
    town_name = town_cfg["name"]
    slug      = town_cfg["slug"]

    add_section_header(doc, "Residential Building Approvals",
                       f"New private residential dwellings approved — {town_name} SA2")
    add_spacer(doc, 8)

    cache = REPO_ROOT / "cache" / "housing" / f"{slug}_qgso.json"
    app_data = {}
    if cache.exists():
        with open(cache) as f:
            d = json.load(f)
        app_data = {yr: v for yr, v in
                    d.get("indicators", {}).get("building_approvals", {}).items() if v}

    build_chart(doc,
        series=[{"label": "Residential approvals", "data": app_data, "color": HEX_PURPLE}],
        format_fn=fmt_int)

    add_spacer(doc, 8)
    _data_notes(doc, [
        ("Source: ", True),
        ("QGSO Regional Database — Building Approvals (ABS 8731.0).", False),
        ("  Boundary: ", True),
        (f"SA2 ({town_name}, ASGS 2021). Private buildings only.", False),
        ("  Aggregation: ", True),
        ("Annual total = sum of 12 monthly values "
         "(historical series Jul 2001–Dec 2018 merged with current series Jan 2019–present).", False),
    ])
    add_footer_line(doc, town_name)


def build_rainfall_page(doc: Document, town_cfg: dict):
    import json
    from common import REPO_ROOT
    town_name = town_cfg["name"]
    slug      = town_cfg["slug"]
    station   = town_cfg.get("bom_station", "")

    add_section_header(doc, "Rainfall",
                       f"Annual total rainfall — {town_name}"
                       + (f" (SILO station {station})" if station else ""))
    add_spacer(doc, 8)

    cache = REPO_ROOT / "cache" / "rainfall" / f"{slug}_bom_rainfall.json"
    rain_data = {}
    if cache.exists():
        with open(cache) as f:
            d = json.load(f)
        rain_data = {yr: v for yr, v in
                     d.get("indicators", {}).get("rainfall", {}).items() if v}

    if rain_data:
        avg = sum(rain_data.values()) / len(rain_data)
        build_chart(doc,
            series=[{"label": "Annual rainfall (mm)", "data": rain_data, "color": HEX_TEAL}],
            format_fn=fmt_mm)
        add_spacer(doc, 4)
        avg_p = doc.add_paragraph()
        avg_p.paragraph_format.space_before = Pt(0)
        avg_p.paragraph_format.space_after  = Pt(4)
        set_run_font(
            avg_p.add_run(f"Historic average ({min(rain_data.keys())}–{max(rain_data.keys())}): "
                          f"{avg:.0f} mm/year"),
            8, italic=True, colour=PURPLE)
    else:
        p = doc.add_paragraph()
        set_run_font(p.add_run("⚠  No rainfall data available."), 9, italic=True, colour=GREY_DARK)

    add_spacer(doc, 4)
    _data_notes(doc, [
        ("Source: ", True),
        ("SILO Patched Point Dataset (Queensland Government / Bureau of Meteorology).", False),
        ("  Station: ", True),
        (f"SILO station {station}." if station else "See towns.toml for station number.", False),
        ("  Note: ", True),
        ("SILO interpolates missing daily observations. "
         "Years with >10% interpolated days are flagged in the pipeline quality log.", False),
    ])
    add_footer_line(doc, town_name)


def build_crime_page(doc: Document, town_cfg: dict):
    import json
    from common import REPO_ROOT
    town_name = town_cfg["name"]
    slug      = town_cfg["slug"]

    add_section_header(doc, "Crime",
                       f"Offences per 1,000 persons per year — {town_name} QPS Division")
    add_spacer(doc, 8)

    cache = REPO_ROOT / "cache" / "crime" / f"{slug}_crime_qps.json"
    indicators = {}
    if cache.exists():
        with open(cache) as f:
            d = json.load(f)
        for key in ["all", "drug", "good_order", "theft", "traffic"]:
            indicators[key] = {yr: v for yr, v in
                               d.get("indicators", {}).get(key, {}).items() if v is not None}

    sub_pages = [
        ("all",        "All offences (per 1,000 persons)",    HEX_PURPLE),
        ("drug",       "Drug offences",                       HEX_PURPLE),
        ("good_order", "Good order offences",                 HEX_TEAL),
        ("theft",      "Theft offences",                      HEX_TEAL),
        ("traffic",    "Traffic offences",                    HEX_PURPLE),
    ]

    for key, label, colour in sub_pages:
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.space_before = Pt(6)
        sub_p.paragraph_format.space_after  = Pt(3)
        set_run_font(sub_p.add_run(label), 9, bold=True, colour=PURPLE)
        build_chart(doc,
            series=[{"label": label, "data": indicators.get(key, {}), "color": colour}],
            format_fn=fmt_rate)
        add_spacer(doc, 4)

    _data_notes(doc, [
        ("Source: ", True),
        ("Queensland Police Service (QPS) open crime data.", False),
        ("  Unit: ", True),
        ("Rates per 100,000 persons converted to per 1,000 persons.", False),
        ("  Geography: ", True),
        ("QPS Division level.", False),
    ])
    add_footer_line(doc, town_name)


# ── Shared data notes helper ──────────────────────────────────────────────────

def _data_notes(doc, items: list[tuple[str, bool]]):
    """items: list of (text, bold) tuples concatenated into one paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    for text, bold in items:
        r = p.add_run(text)
        set_run_font(r, 8, bold=bold, colour=GREY_DARK)